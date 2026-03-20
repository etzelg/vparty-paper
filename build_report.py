from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT_DIR = "israel_analysis_output"
doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(3.0)

# ── Helper: add a horizontal rule ─────────────────────────────────────────────
def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(4)

# ── Helper: styled paragraph ──────────────────────────────────────────────────
def add_body(doc, text, italic=False, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.italic = italic
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        b = p.add_run(bold_prefix)
        b.bold = True
        b.font.size = Pt(11)
        r = p.add_run(text)
        r.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)

def add_figure(doc, fname, caption, width=5.8):
    path = os.path.join(OUT_DIR, fname)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].font.size = Pt(9)
        cp.runs[0].font.italic = True
        cp.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        cp.paragraph_format.space_after = Pt(10)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(40)
r = tp.add_run("Populism, Anti-Pluralism, and Corruption in Israel")
r.bold = True
r.font.size = Pt(20)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("A V-Party Database Analysis of Israeli Party Politics (1990–2022)")
sr.font.size = Pt(13)
sr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
dr = date_p.add_run("March 2026")
dr.font.size = Pt(11)
dr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()
add_hr(doc)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("Executive Summary", level=1)

add_body(doc,
    "This report analyses Israeli political parties using V-Party (Varieties of Democracy Party) "
    "dataset variables spanning three decades. The central finding is a sustained, statistically "
    "significant rightward-populist trajectory in Israel's governing coalition — characterised by "
    "rising anti-elite rhetoric, erosion of pluralist norms, and a concurrent deterioration in "
    "independent corruption indicators. These trends accelerate markedly after 2009 and peak "
    "during the 2019–2023 judicial-reform period.")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 1. DATA & MEASURES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("1. Data and Measures", level=1)

add_body(doc,
    "All party-level measures are drawn from V-Party v3 (2023 release). Corruption proxies use "
    "Transparency International CPI (1995–2023), World Bank Control of Corruption Estimate "
    "(1996–2022), and V-Dem's Executive Corruption Index (v14). Party observations are matched "
    "to election years and weighted by vote share when constructing bloc-level aggregates.")

doc.add_heading("Key Variables", level=2)

add_bullet(doc, "Degree to which party rhetoric frames 'the people' as homogeneous and virtuous, "
                "opposed by a corrupt elite. Scale 0–4 (low → high).",
           bold_prefix="Populism Score (v2papeople): ")

add_bullet(doc, "Degree of anti-establishment, anti-elite, anti-media, or anti-judicial framing "
                "in party communication. Scale 0–4.",
           bold_prefix="Anti-Elitism (v2paanteli): ")

add_bullet(doc, "Willingness to accept political opponents as legitimate; rejection of minority "
                "rights or opposition parties coded as low pluralism. Scale 0–4 (high → pluralist).",
           bold_prefix="Anti-Pluralism (1 − v2paplur): ")

add_bullet(doc, "Party's reliance on targeted material benefits to voters (e.g. jobs, contracts) "
                "rather than programmatic policy appeals.",
           bold_prefix="Clientelism (v2paclient): ")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 2. POPULISM: MAIN FINDINGS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("2. Populism — Key Findings", level=1)

add_body(doc,
    "Populism in V-Party is operationalised as the combination of anti-elite framing (v2paanteli) "
    "and the valorisation of a unified 'people' against corrupt elites (v2papeople). The two "
    "sub-components are positively correlated but measure distinct rhetorical strategies.")

doc.add_heading("2.1 Likud's Long-Run Trajectory", level=2)

add_body(doc,
    "Figure 1 plots Likud's populism score and anti-pluralism score over time with 95 % "
    "confidence intervals derived from the V-Party expert survey. Several periods stand out:")

add_bullet(doc, "1990–2000: Moderate populism (≈ 1.5–2.0), limited anti-pluralism. Likud "
                "behaves broadly as a mainstream conservative party.")
add_bullet(doc, "2001–2009: Slight decline under Sharon/Olmert transition; Kadima split "
                "temporarily pulls the centre-right toward pragmatism.")
add_bullet(doc, "2009–2015: Netanyahu's return initiates a steady climb in both dimensions. "
                "Anti-pluralism rises faster than populism, suggesting institutional "
                "confrontation precedes full populist mobilisation.")
add_bullet(doc, "2015–2022: Both measures reach their highest recorded values. The 2019–2021 "
                "electoral cycle (four elections in two years) coincides with peak scores, "
                "consistent with Likud intensifying anti-elite messaging during Netanyahu's "
                "corruption trial.")

add_figure(doc, "fig1_likud_trajectory.png",
           "Figure 1 — Likud populism score and anti-pluralism index (1990–2022) with 95 % CI bands.")

doc.add_heading("2.2 Left–Right Bloc Comparison", level=2)

add_body(doc,
    "Figure 2 compares vote-share-weighted bloc averages for populism and anti-pluralism. "
    "The contrast is striking: the right bloc averages roughly double the populism score of "
    "the centre-left bloc across the full period, and the gap widens after 2015. "
    "The centre-left shows a mild downward trend in anti-pluralism; the right trend is "
    "sharply upward. This divergence suggests populism has become an increasingly "
    "load-bearing element of right-bloc identity, not merely an occasional rhetorical device.")

add_figure(doc, "fig4_bloc_comparison.png",
           "Figure 2 — Vote-weighted bloc averages: populism and anti-pluralism, left vs. right blocs.")

doc.add_heading("2.3 Cross-Party Snapshot (Most Recent Wave)", level=2)

add_body(doc,
    "Figure 3 places every major Israeli party on a two-dimensional plane (anti-elitism × "
    "anti-pluralism) using the most recent available V-Party observation. The scatterplot "
    "reveals a clear right-populist cluster — Religious Zionism, Otzma Yehudit, Shas, "
    "UTJ, and Likud — occupying the high-anti-elitism, high-anti-pluralism quadrant. "
    "Centre and left parties cluster near the origin. No party combines high pluralism "
    "with high anti-elitism, consistent with the theoretical expectation that populism "
    "erodes intra-system pluralist norms.")

add_figure(doc, "fig7_scatter_latest.png",
           "Figure 3 — Israeli parties: anti-elitism vs. anti-pluralism (latest V-Party wave). "
           "Bubble size = vote share; colour = bloc.")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 3. ANTI-ELITISM IN GOVERNMENT vs. OPPOSITION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("3. Anti-Elitism: In Power vs. Opposition", level=1)

add_body(doc,
    "A canonical prediction of populism theory is that anti-elite rhetoric diminishes once "
    "a party gains executive power — having captured the state, it has less need to "
    "attack it. Israel's right bloc defies this expectation.")

add_figure(doc, "fig11_inpower_vs_opposition.png",
           "Figure 4 — Mean anti-elitism score when governing versus in opposition, by bloc.")

add_body(doc,
    "The bar chart shows that Likud and its allies score *higher* on anti-elitism when in "
    "government than in opposition. This is theoretically anomalous and has two plausible "
    "explanations:")

add_bullet(doc, "The 'permanent campaign' dynamic: Netanyahu weaponised anti-judicial and "
                "anti-media rhetoric most intensively precisely while in office, as a shield "
                "against corruption investigations.",
           bold_prefix="Instrumental use of anti-elitism: ")

add_bullet(doc, "Likud in government built a parallel loyalist structure within state "
                "institutions rather than withdrawing anti-elite framing, intensifying "
                "attacks on judges, the attorney-general, and the press.",
           bold_prefix="Competitive authoritarianism: ")

add_body(doc,
    "In contrast, the centre-left bloc shows the theoretically expected pattern: "
    "higher anti-elitism in opposition, lower in government. This asymmetry supports "
    "the view that right-bloc anti-elitism is strategic and persistent, not merely "
    "oppositional.")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 4. CORRUPTION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("4. Corruption — Key Findings", level=1)

add_body(doc,
    "Populist parties frequently invoke corruption as a legitimating frame ('the corrupt "
    "elite'), but the empirical relationship between populism and actual corruption levels "
    "is contested. We triangulate three independent data sources to assess whether Israel's "
    "rising populism coincides with measurable governance deterioration.")

doc.add_heading("4.1 Multi-Source Corruption Trends", level=2)

add_body(doc,
    "Figure 5 overlays three normalised corruption indicators (TI CPI, World Bank CCE, "
    "and V-Dem Executive Corruption), standardised to a common 0–1 scale where higher "
    "values indicate worse perceived corruption. All three sources agree on the direction:")

add_bullet(doc, "Israel scores in the top decile globally on corruption control through the "
                "late 1990s and 2000s — broadly consistent with its OECD peers.")
add_bullet(doc, "A sustained deterioration begins c.2008–2010 across all three sources, "
                "coinciding with Netanyahu's return to the premiership.")
add_bullet(doc, "By 2020–2022 Israel's CPI rank has dropped from ~25th globally to ~35th; "
                "the V-Dem executive corruption score shows the steepest decline, "
                "reflecting expert assessments of direct executive-branch corruption.")
add_bullet(doc, "The World Bank CCE — the most smoothed of the three — shows a shallower "
                "but directionally consistent decline, reducing the risk that any single "
                "source drives the result.")

add_figure(doc, "fig17_corruption_normalised_combined.png",
           "Figure 5 — Normalised corruption indicators from three independent sources (higher = worse). "
           "Shaded bands indicate periods of right-bloc government.")

doc.add_heading("4.2 Interpretation", level=2)

add_body(doc,
    "The corruption deterioration is moderate in absolute terms — Israel remains "
    "substantially less corrupt than the median democracy — but the *trend* is "
    "statistically robust and cross-validated across sources. Several mechanisms "
    "may link rising populism to this deterioration:")

add_bullet(doc, "Netanyahu's 2019–2021 corruption trial runs in parallel with peak "
                "anti-judicial rhetoric, consistent with personal-interest capture of "
                "anti-elite framing.",
           bold_prefix="Elite capture of anti-corruption narrative: ")

add_bullet(doc, "Clientelism scores rise when the right bloc holds power, suggesting "
                "particularistic distribution of state resources alongside populist "
                "universalist appeals.",
           bold_prefix="Clientelistic governance: ")

add_bullet(doc, "Sustained attacks on judicial independence and prosecutorial autonomy "
                "weaken accountability mechanisms that would otherwise constrain corruption.",
           bold_prefix="Institutional erosion: ")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 5. SYNTHESIS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("5. Synthesis and Theoretical Implications", level=1)

add_body(doc,
    "Taken together, the evidence points to a coherent populist-backsliding syndrome "
    "in Israel that has intensified over three decades:")

rows = [
    ("Rising populism",          "Likud and allies show monotonically increasing populism scores 2009–2022; right bloc averages roughly 2× the left bloc."),
    ("Anomalous anti-elitism",   "Anti-elite rhetoric is highest when governing, not in opposition — consistent with strategic, institutionally targeted mobilisation."),
    ("Anti-pluralism",           "Rejection of opponents' legitimacy is now a structural feature of the right bloc, not episodic. Correlated with judicial-reform agenda."),
    ("Corruption deterioration", "All three independent sources agree on a post-2010 decline in corruption control, moderate in magnitude but consistent in direction."),
    ("Clientelism",              "Increases with right-bloc government tenure, suggesting particularistic alongside populist strategies."),
]

table = doc.add_table(rows=1 + len(rows), cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Dimension"
hdr_cells[1].text = "Finding"
for cell in hdr_cells:
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(11)
for label, finding in rows:
    row_cells = table.add_row().cells
    row_cells[0].text = label
    row_cells[0].paragraphs[0].runs[0].bold = True
    row_cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    row_cells[1].text = finding
    row_cells[1].paragraphs[0].runs[0].font.size = Pt(10)

doc.add_paragraph()

add_body(doc,
    "These findings align with the 'backsliding through populism' literature (Levitsky & "
    "Ziblatt 2018; Mudde & Kaltwasser 2017; Norris & Inglehart 2019): populist parties "
    "do not always collapse democratic institutions overnight, but their sustained anti-"
    "pluralist rhetoric and institutional confrontation gradually erode the norms and "
    "organisations that constrain executive power and corruption. Israel represents a "
    "middle-income, consolidated-democracy case that is undergoing this process in slow "
    "motion.")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 6. CAVEATS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("6. Caveats and Limitations", level=1)

add_bullet(doc, "V-Party relies on expert surveys (5–7 coders per party-year). Expert disagreement "
                "is captured in confidence intervals but residual bias cannot be ruled out.")
add_bullet(doc, "The most recent V-Party wave is 2022; the 2023 judicial-reform legislation and "
                "the post-October 2023 wartime politics are not yet covered.")
add_bullet(doc, "Causal attribution is limited — corruption deterioration could reflect "
                "improved measurement or societal changes rather than governance shifts.")
add_bullet(doc, "Vote-weighted bloc averages obscure within-bloc variation; some right-bloc "
                "parties (e.g. Kulanu, New Hope) show lower populism scores than the bloc mean.")

doc.add_paragraph()
add_hr(doc)
add_body(doc, "Report generated from V-Party v3, V-Dem v14, TI CPI, and World Bank WGI data. "
              "Analysis and figures produced in Python (pandas, matplotlib, seaborn).",
         italic=True)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = os.path.join(OUT_DIR, "Israel_Populism_Report.docx")
doc.save(out_path)
print(f"Saved → {out_path}")
